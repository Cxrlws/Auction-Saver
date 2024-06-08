
import tkinter as tk
from tkinter import messagebox
from docx import Document

class AuctionApp:
    def __init__(self, master):
        self.master = master
        master.title("Registro de Arremates")
        self.create_widgets()
        self.create_table()

    def create_widgets(self):
        self.total_value = 0.0
        self.total_pix = 0.0
        self.total_cartao = 0.0
        self.total_dinheiro = 0.0

        self.label_name = tk.Label(self.master, text="Nome do Item:")
        self.label_name.grid(row=0, column=0)
        self.entry_name = tk.Entry(self.master)
        self.entry_name.grid(row=0, column=1)

        self.label_value = tk.Label(self.master, text="Valor do Item:")
        self.label_value.grid(row=1, column=0)
        self.entry_value = tk.Entry(self.master)
        self.entry_value.grid(row=1, column=1)

        self.label_bidder = tk.Label(self.master, text="Nome do Arrematante:")
        self.label_bidder.grid(row=2, column=0)
        self.entry_bidder = tk.Entry(self.master)
        self.entry_bidder.grid(row=2, column=1)

        self.label_payment = tk.Label(self.master, text="Forma de Pagamento (PIX/Cartão/Dinheiro):")
        self.label_payment.grid(row=3, column=0)
        self.entry_payment = tk.Entry(self.master)
        self.entry_payment.grid(row=3, column=1)

        self.label_paid = tk.Label(self.master, text="Pago? (s/n):")
        self.label_paid.grid(row=4, column=0)
        self.entry_paid = tk.Entry(self.master)
        self.entry_paid.grid(row=4, column=1)

        # Associa a tecla "Enter" para ir para a próxima entrada
        self.entry_name.bind("<Return>", lambda event: self.entry_value.focus_set())
        self.entry_value.bind("<Return>", lambda event: self.entry_bidder.focus_set())
        self.entry_bidder.bind("<Return>", lambda event: self.entry_payment.focus_set())
        self.entry_payment.bind("<Return>", lambda event: self.entry_paid.focus_set())
        self.entry_paid.bind("<Return>", self.add_item)

        self.button_add = tk.Button(self.master, text="Adicionar Item", command=self.add_item)
        self.button_add.grid(row=5, column=0, columnspan=2)

        self.button_create_file = tk.Button(self.master, text="Criar Arquivo e Encerrar", command=self.create_file_and_exit)
        self.button_create_file.grid(row=6, column=0, columnspan=2)

        self.button_quit = tk.Button(self.master, text="Sair", command=self.master.quit)
        self.button_quit.grid(row=7, column=0, columnspan=2)

    def create_table(self):
        self.doc = Document()
        self.doc.add_heading('Registro de Arremates', level=1)
        self.table = self.doc.add_table(rows=1, cols=5)
        hdr_cells = self.table.rows[0].cells
        hdr_cells[0].text = 'Nome do Item'
        hdr_cells[1].text = 'Valor'
        hdr_cells[2].text = 'Nome do Arrematante'
        hdr_cells[3].text = 'Forma de Pagamento'
        hdr_cells[4].text = 'Pago?'

    def add_item(self, event=None):
        try:
            item_name = self.entry_name.get()
            item_value = float(self.entry_value.get())
            bidder_name = self.entry_bidder.get()
            payment_method = self.entry_payment.get().lower()
            if payment_method not in ['pix', 'cartão', 'dinheiro']:
                raise ValueError("Opção de pagamento inválida.")
            paid = self.entry_paid.get().lower()
            if paid not in ['s', 'n']:
                raise ValueError("Opção de pagamento inválida.")

            # Adiciona os detalhes do item na tabela do documento
            row_cells = self.table.add_row().cells
            row_cells[0].text = item_name
            row_cells[1].text = str(item_value)
            row_cells[2].text = bidder_name
            row_cells[3].text = payment_method
            row_cells[4].text = paid.upper()

            self.total_value += item_value
            if payment_method == 'pix':
                self.total_pix += item_value
            elif payment_method == 'cartão':
                self.total_cartao += item_value
            elif payment_method == 'dinheiro':
                self.total_dinheiro += item_value

            self.clear_entries()
            messagebox.showinfo("Sucesso", "Item adicionado.")

        except ValueError as e:
            messagebox.showerror("Erro", e)
            self.clear_entries()

    def create_file_and_exit(self):
        try:
            self.doc.add_paragraph(f"Soma total dos valores: {self.total_value:.2f}")
            self.doc.add_paragraph(f"Valor total (PIX): {self.total_pix:.2f}")
            self.doc.add_paragraph(f"Valor total (Cartão): {self.total_cartao:.2f}")
            self.doc.add_paragraph(f"Valor total (Dinheiro): {self.total_dinheiro:.2f}")

            self.doc.save('arremate.docx')

            messagebox.showinfo("Sucesso", "Arquivo criado com sucesso.")
            self.master.quit()

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao criar o arquivo: {e}")

    def clear_entries(self):
        self.entry_name.delete(0, tk.END)
        self.entry_value.delete(0, tk.END)
        self.entry_bidder.delete(0, tk.END)
        self.entry_payment.delete(0, tk.END)
        self.entry_paid.delete(0, tk.END)
        self.entry_name.focus_set()

def main():
    root = tk.Tk()
    app = AuctionApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
